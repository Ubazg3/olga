from pathlib import Path
import sys

from docx import Document


doc_path = Path(sys.argv[1])
doc = Document(str(doc_path))

print(f"path={doc_path}")
print(f"paragraphs={len(doc.paragraphs)}")
print(f"tables={len(doc.tables)}")

for i, p in enumerate(doc.paragraphs):
    text = p.text.replace("\t", " ").strip()
    if not text:
        continue
    style = p.style.name if p.style is not None else ""
    if (
        text.startswith("1.")
        or text.startswith("2.")
        or text.startswith("3.")
        or text.startswith("4.")
        or text.startswith("5.")
        or text.startswith("6.")
        or text.startswith("7.")
        or text.startswith("8.")
        or text.startswith("9.")
        or text.startswith("10.")
        or text.startswith("11.")
        or text.startswith("12.")
        or text.startswith("13.")
        or text.startswith("14.")
        or "פרק" in text
        or "חלק" in text
        or "מענה למחוון" in text
        or "ניתוח" in text
    ):
        print(f"{i:04d} | {style} | {text[:180]}")
